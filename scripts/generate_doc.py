"""Generate the comprehensive implementation guide Word document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return table


def build_document():
    doc = Document()

    # -- Styles --
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ================================================================
    # TITLE PAGE
    # ================================================================
    doc.add_paragraph("")
    doc.add_paragraph("")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Bayshore QBO Pipeline")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("QuickBooks Online to Power BI\nMulti-Client Data Pipeline")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph("")
    pilot = doc.add_paragraph()
    pilot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pilot.add_run("Pilot Client: Tampa Fence")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 102, 153)

    doc.add_paragraph("")
    doc.add_paragraph("")
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Prepared by: Bayshore Biz Solutions\n").font.size = Pt(11)
    info.add_run("February 2026").font.size = Pt(11)

    doc.add_page_break()

    # ================================================================
    # TABLE OF CONTENTS
    # ================================================================
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Architecture Overview",
        "3. What Was Built (Complete File Reference)",
        "4. Prerequisites & Environment Setup",
        "5. Step-by-Step: Local Development Setup",
        "6. Step-by-Step: Onboard Tampa Fence (Pilot)",
        "7. Step-by-Step: Build the Power BI Report",
        "8. Step-by-Step: Deploy to Azure (Production)",
        "9. Step-by-Step: Add a New Client",
        "10. Ongoing Maintenance & Monitoring",
        "11. Power BI Report Design (8 Pages)",
        "12. DAX Measures Reference",
        "13. Cost Breakdown",
        "14. Troubleshooting",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ================================================================
    # 1. EXECUTIVE SUMMARY
    # ================================================================
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "This document describes a complete data pipeline that automatically extracts "
        "financial data from QuickBooks Online (QBO), transforms it into a structured "
        "star-schema database, and visualizes it in Power BI. The system is designed to "
        "serve multiple small business clients through Bayshore Biz Solutions, with "
        "Tampa Fence as the pilot deployment."
    )
    doc.add_paragraph(
        "The pipeline replaces or augments the weekly and monthly financial questions "
        "that small businesses typically ask their CPA, including:"
    )
    bullets = [
        "How is my revenue trending? Am I up or down vs last month/year?",
        "What does my Profit & Loss look like?",
        "What is my current balance sheet position?",
        "Who owes me money and how overdue is it? (AR Aging)",
        "Who do I owe money to? (AP Aging)",
        "What are my biggest expense categories?",
        "What is my cash position and cash flow trend?",
        "How are my estimates converting to invoices?",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("Key Design Decisions", level=2)
    doc.add_paragraph(
        "Azure Functions + Azure SQL Serverless was chosen over Databricks because the "
        "data volumes for small businesses (1,000-10,000 transactions/month) do not justify "
        "Spark. The estimated monthly cost is $15-35 for infrastructure (excluding Power BI "
        "Pro licenses at $10/user/month)."
    )

    doc.add_page_break()

    # ================================================================
    # 2. ARCHITECTURE OVERVIEW
    # ================================================================
    doc.add_heading("2. Architecture Overview", level=1)

    doc.add_heading("Data Flow", level=2)
    flow_text = (
        "QuickBooks Online API (per client)\n"
        "    \u2193\n"
        "Azure Function (daily timer trigger, 6:00 AM ET)\n"
        "    \u2193 Extract: REST API calls per table per client\n"
        "    \u2193 Transform: flatten nested JSON, normalize line items\n"
        "    \u2193 Load: upsert into Azure SQL star schema\n"
        "    \u2193\n"
        "Power BI Service (scheduled refresh, up to 8x/day)\n"
        "    \u2193\n"
        "Client views their dashboard (filtered by Row-Level Security)"
    )
    p = doc.add_paragraph()
    run = p.add_run(flow_text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)

    doc.add_heading("Infrastructure Components", level=2)
    add_styled_table(doc,
        ["Component", "Service", "Monthly Cost"],
        [
            ["ETL Orchestration", "Azure Functions (Consumption Plan)", "~$0 (free tier)"],
            ["Database", "Azure SQL Serverless (1 vCore, auto-pause)", "~$5-15"],
            ["Token Storage", "Azure Key Vault", "~$0"],
            ["Raw JSON Backup", "Azure Blob Storage", "<$1"],
            ["Reporting", "Power BI Pro per client", "$10/user"],
            ["TOTAL", "", "$15-35 + PBI licenses"],
        ]
    )

    doc.add_heading("Multi-Client Isolation", level=2)
    doc.add_paragraph(
        "Every table in the database has a client_id column. Power BI uses Row-Level "
        "Security (RLS) to ensure each client only sees their own data. Each client gets "
        "a separate Power BI workspace."
    )

    doc.add_heading("QBO Data Sources", level=2)
    doc.add_paragraph("The pipeline extracts 24 entity tables and 9 report endpoints from the QBO API:")
    add_styled_table(doc,
        ["Category", "Tables/Endpoints"],
        [
            ["Core Transactions", "Invoice, Payment, Bill, Purchase, Estimate, SalesReceipt"],
            ["Supporting Transactions", "BillPayment, Deposit, CreditMemo, RefundReceipt, JournalEntry, Transfer"],
            ["Dimensions", "Customer, Vendor, Employee, Item, Account, Department, Class"],
            ["Reference", "TaxCode, TaxRate, Term, PaymentMethod, CompanyInfo"],
            ["Reports (GET)", "ProfitAndLoss (Accrual & Cash), BalanceSheet, CashFlow, AgedReceivables, AgedReceivableDetail, AgedPayables, AgedPayableDetail, TrialBalance"],
        ]
    )

    doc.add_page_break()

    # ================================================================
    # 3. WHAT WAS BUILT
    # ================================================================
    doc.add_heading("3. What Was Built (Complete File Reference)", level=1)
    doc.add_paragraph(
        "All code lives in C:\\Users\\sbien\\bayshore-qbo-pipeline\\. "
        "No existing files were modified."
    )

    files = [
        ("config/settings.py", "Pydantic Settings class. Reads environment variables for QBO credentials, Azure connection strings, database backend (sqlite vs azure_sql), and token storage mode (local vs keyvault)."),
        ("config/tables.py", "Defines all 24 QBO entity tables, 9 report endpoints with default parameters, the line-item detail type mappings, and the entity-to-schema-table mapping."),
        ("config/clients.json", "JSON registry of onboarded clients. Each entry has client_id, client_name, qbo_realm_id, industry, onboarded_date, and is_active flag. Created by onboard_client.py."),
        ("auth/oauth_manager.py", "OAuth2 token lifecycle management. Stores tokens in Azure Key Vault (production) or local JSON files (development). Auto-refreshes access tokens before expiry (60-min lifetime). Warns when refresh tokens approach 101-day expiry."),
        ("auth/qbo_auth_flow.py", "One-time OAuth2 authorization code flow. Starts a local HTTP server, opens the Intuit authorization page in the browser, captures the callback code, exchanges it for tokens, and stores them."),
        ("extract/qbo_client.py", "Core QBO REST API client. Handles pagination (STARTPOSITION/MAXRESULTS), rate limiting (500 req/min), retry with exponential backoff on 429/5xx errors, and both POST (entity query) and GET (report) endpoints."),
        ("extract/entity_extractor.py", "Iterates through all 24 entity tables and calls qbo_client.query_all() for each. Returns a dict of {table_name: [records...]}."),
        ("extract/report_extractor.py", "Fetches QBO report endpoints (P&L, Balance Sheet, etc.) with date parameters. Includes a recursive flattener for the nested QBO report JSON structure (Sections > Rows > ColData)."),
        ("transform/flatteners.py", "The most complex module. Flattens nested Line arrays from Invoice (SalesItemLineDetail), Bill (AccountBasedExpenseLineDetail + ItemBasedExpenseLineDetail), Purchase, Payment (LineEx NameValue pairs), and Estimate into one-row-per-line-item DataFrames."),
        ("transform/schema_mapper.py", "Maps raw QBO API field names to star-schema column names for each dimension and fact table. Handles nested refs (CustomerRef.value, AccountRef.value, etc.). Registry pattern allows map_to_schema(table, records) dispatch."),
        ("transform/data_quality.py", "Type enforcement (cast to float, bool, date, int), deduplication by key columns, and date_key generation (YYYYMMDD integer for dim_date joins). Includes predefined type maps for each table."),
        ("load/sql_loader.py", "Loads DataFrames into the database. SQLite backend uses INSERT OR REPLACE. Azure SQL backend uses a staging table + MERGE (upsert) pattern. Automatically adds client_id to every row. Also handles executing raw .sql files."),
        ("load/raw_archiver.py", "Archives raw QBO API JSON responses. Local dev saves to data/raw_archive/{client}/{table}/{timestamp}.json. Production uploads to Azure Blob Storage."),
        ("orchestrator/pipeline.py", "Main ETL orchestrator. For each client: initializes OAuth, extracts all entities, transforms and loads dimension tables, flattens and loads line-item fact tables, extracts and loads reports. CLI entry point for manual runs."),
        ("function_app.py", "Azure Functions entry point. Timer trigger fires daily at 11:00 UTC (6:00 AM ET). Calls run_pipeline_all_clients() which iterates through clients.json."),
        ("host.json", "Azure Functions host configuration. Specifies runtime version 2.0 and extension bundle."),
        ("scripts/create_schema.sql", "DDL for all star-schema tables. 13 dimension tables + 14 fact tables + 1 security mapping table. Uses SQLite-compatible syntax (CREATE TABLE IF NOT EXISTS, TEXT/REAL/INTEGER types). Every table has client_id in its primary key."),
        ("scripts/populate_dim_date.sql", "Populates dim_date with every date from 2020-01-01 through 2030-12-31 using a recursive CTE. Includes day/week/month/quarter/year numbers, names, fiscal calendar fields, and weekend flag."),
        ("scripts/onboard_client.py", "CLI tool for adding a new client. Runs the OAuth flow, stores tokens, registers the client in clients.json, and optionally initializes the database schema (--init-db flag)."),
        ("scripts/backfill.py", "One-time historical data load. Pulls entity tables (which contain all historical data) and pulls report endpoints year-by-year for trend analysis."),
        ("powerbi/dax_measures.dax", "All 45+ DAX measures organized by display folder: Revenue, AR, AP, Expenses, Profitability, Cash Flow, Estimates Pipeline, and Tax."),
        ("powerbi/power_query_sqlite.m", "Power Query M scripts for connecting Power BI Desktop to the local SQLite database (development). One query per table with type transformations."),
        ("powerbi/power_query_azure_sql.m", "Power Query M scripts for connecting Power BI Desktop to Azure SQL (production). One query per table."),
        ("requirements.txt", "Python dependencies: requests, pandas, pydantic-settings, pyodbc, azure-functions, azure-identity, azure-keyvault-secrets, azure-storage-blob, python-dotenv."),
        (".env.example", "Template for environment variables. QBO credentials, Azure connection strings, database backend toggle."),
        (".gitignore", "Ignores .env, tokens, CSV data, __pycache__, .pbix files, and clients.json (contains realm IDs)."),
    ]

    for filename, description in files:
        p = doc.add_paragraph()
        run = p.add_run(filename)
        run.bold = True
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        p.add_run(f"\n{description}")
        p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ================================================================
    # 4. PREREQUISITES
    # ================================================================
    doc.add_heading("4. Prerequisites & Environment Setup", level=1)

    doc.add_heading("Software Requirements", level=2)
    prereqs = [
        ("Python 3.12+", "Already installed on your machine"),
        ("Power BI Desktop", "Free download from Microsoft; already installed"),
        ("Git", "For version control"),
        ("Azure CLI", "For deploying to Azure; install from https://aka.ms/installazurecli"),
        ("ODBC Driver for SQL Server", "Version 18; needed for Azure SQL connection via pyodbc"),
        ("SQLite ODBC Driver", "(Optional) Only if connecting Power BI directly to SQLite for dev"),
    ]
    for name, note in prereqs:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(name)
        run.bold = True
        p.add_run(f" - {note}")

    doc.add_heading("Accounts Required", level=2)
    accounts = [
        ("Intuit Developer Account", "https://developer.intuit.com - Register an OAuth app for production. You may already have this for Tampa Fence."),
        ("Azure Subscription", "https://portal.azure.com - Pay-as-you-go subscription. Cost will be $15-35/month."),
        ("Power BI Pro License", "$10/user/month per client viewer. Needed for publishing and sharing reports."),
    ]
    for name, note in accounts:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(name)
        run.bold = True
        p.add_run(f" - {note}")

    doc.add_page_break()

    # ================================================================
    # 5. LOCAL DEV SETUP
    # ================================================================
    doc.add_heading("5. Step-by-Step: Local Development Setup", level=1)

    steps = [
        (
            "Open a terminal and navigate to the project",
            "cd C:\\Users\\sbien\\bayshore-qbo-pipeline"
        ),
        (
            "Create and activate a Python virtual environment",
            "python -m venv .venv\n.venv\\Scripts\\activate"
        ),
        (
            "Install Python dependencies",
            "pip install -r requirements.txt"
        ),
        (
            "Create your .env file from the template",
            "copy .env.example .env\n\nThen open .env in a text editor and fill in:\n"
            "  QBO_CLIENT_ID=<your Intuit app client ID>\n"
            "  QBO_CLIENT_SECRET=<your Intuit app client secret>\n"
            "  QBO_REDIRECT_URI=http://localhost:8080/callback\n\n"
            "For local development, leave the Azure settings empty and ensure:\n"
            "  TOKEN_STORAGE=local\n"
            "  DB_BACKEND=sqlite\n"
            "  LOCAL_DB_PATH=data/dev.db"
        ),
        (
            "Verify the setup by checking that Python can import the modules",
            "python -c \"from config.settings import Settings; print(Settings())\""
        ),
    ]

    for i, (title, cmd) in enumerate(steps, 1):
        doc.add_heading(f"Step {i}: {title}", level=2)
        p = doc.add_paragraph()
        run = p.add_run(cmd)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    doc.add_page_break()

    # ================================================================
    # 6. ONBOARD TAMPA FENCE
    # ================================================================
    doc.add_heading("6. Step-by-Step: Onboard Tampa Fence (Pilot)", level=1)

    doc.add_heading("Step 1: Register your QBO OAuth App (if not already done)", level=2)
    doc.add_paragraph(
        "Go to https://developer.intuit.com and sign in. Navigate to My Apps. "
        "If you already have a production app for Tampa Fence, note the Client ID and "
        "Client Secret. If not, create a new app:\n"
        "  - App name: Bayshore QBO Pipeline\n"
        "  - Scopes: com.intuit.quickbooks.accounting\n"
        "  - Redirect URI: http://localhost:8080/callback\n"
        "  - Environment: Production (or Sandbox for testing)\n\n"
        "Copy the Client ID and Client Secret into your .env file."
    )

    doc.add_heading("Step 2: Initialize the database and onboard the client", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "python scripts/onboard_client.py \\\n"
        "    --client-id tampa_fence \\\n"
        "    --name \"Tampa Fence\" \\\n"
        "    --industry \"Fencing / Construction\" \\\n"
        "    --init-db"
    )
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    doc.add_paragraph(
        "This will:\n"
        "  1. Create the SQLite database with all star-schema tables\n"
        "  2. Populate dim_date (2020-2030)\n"
        "  3. Open your browser to Intuit's authorization page\n"
        "  4. Tampa Fence's owner logs into their QuickBooks and clicks 'Authorize'\n"
        "  5. The callback stores access + refresh tokens locally\n"
        "  6. The client is registered in config/clients.json"
    )

    doc.add_heading("Step 3: Run the initial data pull", level=2)
    p = doc.add_paragraph()
    run = p.add_run("python -m orchestrator.pipeline --client-id tampa_fence")
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    doc.add_paragraph(
        "This extracts all 24 entity tables and 9 reports from Tampa Fence's QBO, "
        "transforms the data, and loads it into the local SQLite database at data/dev.db. "
        "You should see log output like:\n"
        "  Extracting entity: Customer... 15 records\n"
        "  Extracting entity: Invoice... 32 records\n"
        "  etc."
    )

    doc.add_heading("Step 4: Run the historical backfill (optional)", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "python scripts/backfill.py --client-id tampa_fence --start-year 2023"
    )
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    doc.add_paragraph(
        "This pulls P&L, Balance Sheet, and Cash Flow reports for 2023, 2024, and 2025, "
        "giving you historical trend data in Power BI."
    )

    doc.add_page_break()

    # ================================================================
    # 7. BUILD POWER BI REPORT
    # ================================================================
    doc.add_heading("7. Step-by-Step: Build the Power BI Report", level=1)

    doc.add_heading("Step 1: Connect Power BI to the database", level=2)
    doc.add_paragraph(
        "For local development with SQLite:\n"
        "  1. Install the SQLite ODBC driver from http://www.ch-werner.de/sqliteodbc/\n"
        "  2. Open ODBC Data Source Administrator (64-bit)\n"
        "  3. Add a System DSN: Driver = SQLite3 ODBC Driver, Database = full path to data/dev.db\n"
        "  4. In Power BI Desktop: Home > Get Data > ODBC\n"
        "  5. Select your SQLite DSN and load all tables\n\n"
        "For production with Azure SQL:\n"
        "  1. In Power BI Desktop: Home > Get Data > Azure SQL Database\n"
        "  2. Server: bayshore-qbo-sql.database.windows.net\n"
        "  3. Database: qbo-warehouse\n"
        "  4. Authentication: Microsoft Account or Azure AD\n"
        "  5. Load all tables\n\n"
        "Reference Power Query M scripts are in the powerbi/ folder."
    )

    doc.add_heading("Step 2: Set up the data model relationships", level=2)
    doc.add_paragraph(
        "In Power BI Desktop, go to Model view and create these relationships "
        "(all are Many-to-One, single direction):"
    )
    add_styled_table(doc,
        ["From Table (Many)", "From Column", "To Table (One)", "To Column", "Active"],
        [
            ["fact_invoice", "txn_date_key", "dim_date", "date_key", "Yes"],
            ["fact_invoice", "due_date_key", "dim_date", "date_key", "No (inactive)"],
            ["fact_invoice", "customer_id", "dim_customer", "customer_id", "Yes"],
            ["fact_invoice_line", "invoice_id", "fact_invoice", "invoice_id", "Yes"],
            ["fact_invoice_line", "item_id", "dim_item", "item_id", "Yes"],
            ["fact_invoice_line", "account_id", "dim_account", "account_id", "Yes"],
            ["fact_payment", "txn_date_key", "dim_date", "date_key", "Yes"],
            ["fact_payment", "customer_id", "dim_customer", "customer_id", "Yes"],
            ["fact_bill", "txn_date_key", "dim_date", "date_key", "Yes"],
            ["fact_bill", "vendor_id", "dim_vendor", "vendor_id", "Yes"],
            ["fact_bill_line", "bill_id", "fact_bill", "bill_id", "Yes"],
            ["fact_bill_line", "account_id", "dim_account", "account_id", "Yes"],
            ["fact_purchase", "txn_date_key", "dim_date", "date_key", "Yes"],
            ["fact_purchase", "vendor_id", "dim_vendor", "vendor_id", "Yes"],
            ["fact_estimate", "txn_date_key", "dim_date", "date_key", "Yes"],
            ["fact_estimate", "customer_id", "dim_customer", "customer_id", "Yes"],
        ]
    )
    doc.add_paragraph(
        "\nImportant: Mark dim_date as the Date Table (right-click dim_date > "
        "Mark as Date Table > select full_date column)."
    )

    doc.add_heading("Step 3: Create DAX measures", level=2)
    doc.add_paragraph(
        "Open the file powerbi/dax_measures.dax. For each measure:\n"
        "  1. In Power BI Desktop, click on the fact_invoice table (or create a Measures table)\n"
        "  2. Go to Modeling > New Measure\n"
        "  3. Paste the DAX formula\n"
        "  4. Set the Display Folder in Properties (e.g., 'Revenue', 'AR', 'Profitability')\n\n"
        "There are 45+ measures organized into 8 display folders. The key measures for "
        "Tampa Fence are:\n"
        "  - Total Revenue, Revenue YTD, Revenue MoM Change\n"
        "  - Total AR Outstanding, AR aging buckets, DSO\n"
        "  - Total AP Outstanding, AP aging buckets\n"
        "  - Net Income, Gross Margin %, Net Margin %\n"
        "  - Cash Position, Net Cash Flow\n"
        "  - Open Estimates Value, Estimate to Invoice Rate"
    )

    doc.add_heading("Step 4: Build the report pages", level=2)
    doc.add_paragraph("See Section 11 for detailed page-by-page layout specifications.")

    doc.add_heading("Step 5: Save as a template (.pbit)", level=2)
    doc.add_paragraph(
        "Once the report is built:\n"
        "  1. File > Save As > Power BI Template (.pbit)\n"
        "  2. Save as 'Bayshore_QBO_Template.pbit'\n"
        "  3. This template can be reused for every new client\n"
        "  4. When opening the template, Power BI will prompt for connection parameters"
    )

    doc.add_page_break()

    # ================================================================
    # 8. DEPLOY TO AZURE
    # ================================================================
    doc.add_heading("8. Step-by-Step: Deploy to Azure (Production)", level=1)

    doc.add_heading("Step 1: Create Azure resources", level=2)
    cmds = (
        "# Login to Azure\n"
        "az login\n\n"
        "# Create resource group\n"
        "az group create --name rg-bayshore-qbo-pipeline --location eastus\n\n"
        "# Create Azure SQL Server\n"
        "az sql server create \\\n"
        "    --name bayshore-qbo-sql \\\n"
        "    --resource-group rg-bayshore-qbo-pipeline \\\n"
        "    --location eastus \\\n"
        "    --admin-user bayshore_admin \\\n"
        "    --admin-password <CHOOSE_STRONG_PASSWORD>\n\n"
        "# Create Azure SQL Database (Serverless, auto-pause after 60 min)\n"
        "az sql db create \\\n"
        "    --resource-group rg-bayshore-qbo-pipeline \\\n"
        "    --server bayshore-qbo-sql \\\n"
        "    --name qbo-warehouse \\\n"
        "    --compute-model Serverless \\\n"
        "    --edition GeneralPurpose \\\n"
        "    --family Gen5 \\\n"
        "    --min-capacity 0.5 \\\n"
        "    --capacity 1 \\\n"
        "    --auto-pause-delay 60\n\n"
        "# Create Key Vault\n"
        "az keyvault create \\\n"
        "    --name kv-bayshore-qbo \\\n"
        "    --resource-group rg-bayshore-qbo-pipeline \\\n"
        "    --location eastus\n\n"
        "# Create Storage Account\n"
        "az storage account create \\\n"
        "    --name stbayshoreetl \\\n"
        "    --resource-group rg-bayshore-qbo-pipeline \\\n"
        "    --location eastus \\\n"
        "    --sku Standard_LRS\n\n"
        "# Create Function App\n"
        "az functionapp create \\\n"
        "    --name func-bayshore-qbo-etl \\\n"
        "    --resource-group rg-bayshore-qbo-pipeline \\\n"
        "    --consumption-plan-location eastus \\\n"
        "    --runtime python \\\n"
        "    --runtime-version 3.12 \\\n"
        "    --functions-version 4 \\\n"
        "    --storage-account stbayshoreetl \\\n"
        "    --os-type linux"
    )
    p = doc.add_paragraph()
    run = p.add_run(cmds)
    run.font.name = "Consolas"
    run.font.size = Pt(8)

    doc.add_heading("Step 2: Configure the Function App settings", level=2)
    doc.add_paragraph(
        "Set the environment variables in the Function App:\n"
        "  az functionapp config appsettings set \\\n"
        "      --name func-bayshore-qbo-etl \\\n"
        "      --resource-group rg-bayshore-qbo-pipeline \\\n"
        "      --settings \\\n"
        "          QBO_CLIENT_ID=<your_id> \\\n"
        "          QBO_CLIENT_SECRET=<your_secret> \\\n"
        "          AZURE_KEY_VAULT_URL=https://kv-bayshore-qbo.vault.azure.net/ \\\n"
        "          AZURE_SQL_CONNECTION_STRING='<connection_string>' \\\n"
        "          AZURE_STORAGE_CONNECTION_STRING='<storage_conn_string>' \\\n"
        "          TOKEN_STORAGE=keyvault \\\n"
        "          DB_BACKEND=azure_sql"
    )

    doc.add_heading("Step 3: Run the schema DDL on Azure SQL", level=2)
    doc.add_paragraph(
        "Connect to Azure SQL using Azure Data Studio or SSMS and run the contents of "
        "scripts/create_schema.sql and scripts/populate_dim_date.sql. Note: for Azure SQL, "
        "you may need to convert the SQLite-compatible syntax to T-SQL (replace TEXT with "
        "NVARCHAR(MAX), REAL with DECIMAL(18,2), INTEGER with INT)."
    )

    doc.add_heading("Step 4: Deploy the Function App", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "cd C:\\Users\\sbien\\bayshore-qbo-pipeline\n"
        "func azure functionapp publish func-bayshore-qbo-etl --python"
    )
    run.font.name = "Consolas"
    run.font.size = Pt(9)

    doc.add_heading("Step 5: Migrate tokens to Key Vault", level=2)
    doc.add_paragraph(
        "The local token files need to be migrated to Azure Key Vault. You can do this "
        "by reading the local token JSON and storing it as a Key Vault secret:\n"
        "  az keyvault secret set \\\n"
        "      --vault-name kv-bayshore-qbo \\\n"
        "      --name tampa-fence-token-data \\\n"
        "      --file data/tokens/tampa_fence.json"
    )

    doc.add_heading("Step 6: Publish Power BI to the service", level=2)
    doc.add_paragraph(
        "  1. In Power BI Desktop, update the data source to Azure SQL\n"
        "  2. File > Publish > select the Tampa Fence workspace\n"
        "  3. In Power BI Service, configure scheduled refresh (Settings > Datasets)\n"
        "  4. Set up Row-Level Security roles in the dataset settings"
    )

    doc.add_page_break()

    # ================================================================
    # 9. ADD A NEW CLIENT
    # ================================================================
    doc.add_heading("9. Step-by-Step: Add a New Client", level=1)
    doc.add_paragraph(
        "Once the pipeline is running for Tampa Fence, adding a new client takes ~15 minutes:"
    )
    new_client_steps = [
        "Have the client's QuickBooks owner available (they need to authorize access)",
        "Run: python scripts/onboard_client.py --client-id <new_id> --name \"<Company Name>\"",
        "The client logs into QuickBooks in the browser and clicks Authorize",
        "Run: python -m orchestrator.pipeline --client-id <new_id>",
        "Run: python scripts/backfill.py --client-id <new_id> --start-year 2023",
        "Open the .pbit template in Power BI, it loads the new client's data",
        "Publish to a new Power BI workspace for this client",
        "Add the client's email to security_user_client_map for RLS",
    ]
    for i, step in enumerate(new_client_steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_page_break()

    # ================================================================
    # 10. MAINTENANCE
    # ================================================================
    doc.add_heading("10. Ongoing Maintenance & Monitoring", level=1)

    doc.add_heading("Daily (Automated)", level=2)
    doc.add_paragraph(
        "The Azure Function runs daily at 6:00 AM ET. It automatically:\n"
        "  - Refreshes OAuth access tokens (60-min expiry)\n"
        "  - Pulls all entity tables and reports from QBO\n"
        "  - Upserts data into Azure SQL\n"
        "  - Archives raw JSON to Blob Storage\n"
        "  - Power BI scheduled refresh picks up the new data"
    )

    doc.add_heading("Weekly (Manual Check)", level=2)
    weekly = [
        "Check Azure Function execution logs in the Azure Portal for any failures",
        "Verify Power BI dataset refresh succeeded (Power BI Service > Dataset > Refresh History)",
        "Spot-check one KPI against QBO's built-in reports (e.g., compare Total Revenue YTD)",
    ]
    for item in weekly:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Every 90 Days (Critical)", level=2)
    doc.add_paragraph(
        "QBO refresh tokens expire after 101 days. The pipeline automatically refreshes "
        "them with each daily run, but if a client's pipeline fails for 100+ days, you "
        "must re-authorize:\n"
        "  python -m auth.qbo_auth_flow --client-id <client_id>\n\n"
        "The pipeline logs warnings when a refresh token is older than 90 days."
    )

    doc.add_heading("Monthly", level=2)
    monthly = [
        "Review Azure costs in the Azure Portal > Cost Management",
        "Check Azure SQL storage usage (should stay well under 1 GB for small clients)",
        "Review raw JSON archives in Blob Storage (can set lifecycle policy to delete after 90 days)",
    ]
    for item in monthly:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # ================================================================
    # 11. REPORT DESIGN
    # ================================================================
    doc.add_heading("11. Power BI Report Design (8 Pages)", level=1)
    doc.add_paragraph(
        "Each page is designed to answer specific CPA questions. All pages should include "
        "the Tampa Fence (or Bayshore Biz Solutions) logo in the top-left, a page navigation "
        "strip at the top, and synced date/client slicers."
    )

    pages = [
        (
            "Page 1: Executive Summary",
            "At-a-glance business health dashboard",
            [
                "4 KPI Cards across the top: Revenue YTD (with YoY % arrow), Net Income YTD (green/red), Total AR Outstanding (with DSO below), Cash Position",
                "Line Chart: Monthly Revenue vs Expenses (trailing 12 months)",
                "Donut Chart: Revenue by Service/Product Category (from dim_item via fact_invoice_line)",
                "Clustered Bar Chart: Top 5 Customers by Revenue",
                "Gauge: Gross Margin % with target line at 40%",
                "Slicer: Date range (month picker), Client selector (hidden in single-client mode)",
            ]
        ),
        (
            "Page 2: Profit & Loss",
            "Traditional income statement view",
            [
                "Matrix Visual (full-width): Rows = Account hierarchy (FullyQualifiedName, Classification filter to Revenue/Expense), Columns = Months from dim_date, Values = sum of amounts",
                "Waterfall Chart (right panel): Revenue > COGS > Gross Profit > Operating Expenses > Net Income",
                "KPI Cards: Total Revenue, Gross Profit, Net Income, Gross Margin %",
                "Slicer: Year selector, Monthly/Quarterly/Annual toggle",
                "Conditional formatting: Red font for negative values, data bars on amounts",
            ]
        ),
        (
            "Page 3: Balance Sheet",
            "Point-in-time financial position",
            [
                "Left Matrix: Assets section (dim_account filtered to Classification = Asset), showing CurrentBalance. Subtotals: Current Assets, Fixed Assets, Total Assets",
                "Right Matrix: Liabilities + Equity section (Classification = Liability or Equity). Subtotals: Current Liabilities, Long-Term Liabilities, Total Liabilities, Equity, Total L+E",
                "Validation Card at bottom: Total Assets = Total L+E (should always balance)",
                "Slicer: As-of date (single date picker)",
            ]
        ),
        (
            "Page 4: Cash Flow",
            "Cash movement and trends",
            [
                "4 KPI Cards: Opening Balance, Cash In (payments), Cash Out (purchases), Closing Balance",
                "Waterfall Chart: Opening > Customer Payments > Vendor Payments > Other > Closing",
                "Line Chart: Daily or weekly cash balance trend",
                "Stacked Bar: Cash inflows by customer",
                "Stacked Bar: Cash outflows by expense category",
            ]
        ),
        (
            "Page 5: AR Aging",
            "Collections management for Tampa Fence",
            [
                "5 KPI Cards (color-coded green to red): Not Yet Due, 1-30 Days, 31-60 Days, 61-90 Days, 90+ Days",
                "Stacked Bar Chart: AR Aging by Customer (each bar has aging bucket segments)",
                "Detail Table: Open invoices with Customer, Invoice#, Date, Due Date, Amount, Balance, Days Outstanding (conditional formatting on aging)",
                "Line Chart: DSO trend over 12 months",
                "Gauge: Collection Rate with 95% target",
                "Interaction: clicking a customer filters the invoice table",
            ]
        ),
        (
            "Page 6: AP Aging",
            "Payables management",
            [
                "4 KPI Cards: Current, 1-30 Days, 31-60 Days, 60+ Days",
                "Stacked Bar Chart: AP Aging by Vendor",
                "Detail Table: Open bills with Vendor, Bill Date, Due Date, Amount, Balance, Days Outstanding",
                "Line Chart: DPO trend",
                "Pie Chart: AP by expense category (from dim_account via fact_bill_line)",
            ]
        ),
        (
            "Page 7: Revenue Deep Dive",
            "Revenue analytics for Tampa Fence",
            [
                "Line + Column Combo Chart: Monthly Revenue (columns) with cumulative line",
                "Small Multiples: Revenue by top 5 product/service categories",
                "Table: Customer ranking with Revenue, Invoice Count, Avg Invoice Value, and sparkline trend",
                "Scatter Plot: Customer Revenue (X) vs Invoice Count (Y), bubble size = avg amount",
                "KPI Cards: Average Invoice Amount, Revenue per Customer, Invoice Count",
            ]
        ),
        (
            "Page 8: Expense Analysis",
            "Cost control and vendor analysis",
            [
                "Treemap: Expenses by category (account_type from dim_account), sized by amount",
                "Line Chart: Monthly expense trend with prior year comparison",
                "Stacked Area Chart: Expense composition over time (top 5 categories + Other)",
                "Table: Top 10 vendors by spend with MoM change column",
                "Bar Chart: Vendor concentration (top 5 vendors as % of total spend)",
                "KPI Cards: Total Expenses, Expense Ratio, Largest Expense Category",
            ]
        ),
    ]

    for title, purpose, visuals in pages:
        doc.add_heading(title, level=2)
        p = doc.add_paragraph()
        run = p.add_run(f"Purpose: {purpose}")
        run.italic = True
        for v in visuals:
            doc.add_paragraph(v, style="List Bullet")
        doc.add_paragraph("")  # spacing

    doc.add_page_break()

    # ================================================================
    # 12. DAX REFERENCE
    # ================================================================
    doc.add_heading("12. DAX Measures Reference", level=1)
    doc.add_paragraph(
        "All DAX measures are in the file powerbi/dax_measures.dax. Below is a summary "
        "organized by display folder. Each measure can be created in Power BI via "
        "Modeling > New Measure."
    )

    dax_groups = [
        ("Revenue (11 measures)", [
            "Total Revenue, Revenue MTD, Revenue QTD, Revenue YTD",
            "Revenue Prior Month, Revenue MoM Change",
            "Revenue Prior Year, Revenue YoY Change",
            "Average Invoice Amount, Invoice Count, Revenue per Customer",
        ]),
        ("Accounts Receivable (9 measures)", [
            "Total AR Outstanding, AR Not Yet Due",
            "AR 1-30 Days, AR 31-60 Days, AR 61-90 Days, AR Over 90 Days",
            "Days Sales Outstanding, Open Invoice Count, Collection Rate",
        ]),
        ("Accounts Payable (7 measures)", [
            "Total AP Outstanding, AP Current",
            "AP 1-30 Days, AP 31-60 Days, AP Over 60 Days",
            "Days Payable Outstanding, Open Bill Count",
        ]),
        ("Expenses (6 measures)", [
            "Total Expenses from Bills, Total Expenses from Purchases, Total Expenses",
            "Expenses MTD, Expenses YTD, Expense Ratio",
        ]),
        ("Profitability (5 measures)", [
            "Net Income, Net Income YTD, Net Margin Pct",
            "Gross Profit, Gross Margin Pct",
        ]),
        ("Cash Flow (4 measures)", [
            "Cash Inflow, Cash Outflow, Net Cash Flow, Cash Position",
        ]),
        ("Estimates Pipeline (4 measures)", [
            "Open Estimates Value, Estimate to Invoice Rate",
            "Estimate Count, Average Estimate Value",
        ]),
        ("Tax (2 measures)", [
            "Total Tax Collected, Tax Collected YTD",
        ]),
    ]

    for group_name, measures in dax_groups:
        doc.add_heading(group_name, level=3)
        for m in measures:
            doc.add_paragraph(m, style="List Bullet")

    doc.add_page_break()

    # ================================================================
    # 13. COST BREAKDOWN
    # ================================================================
    doc.add_heading("13. Cost Breakdown", level=1)

    doc.add_heading("Monthly Infrastructure Costs (1-10 clients)", level=2)
    add_styled_table(doc,
        ["Service", "Tier", "Estimated Monthly Cost"],
        [
            ["Azure Functions", "Consumption (free tier: 1M executions)", "$0"],
            ["Azure SQL Database", "Serverless, 1 vCore, auto-pause 60min", "$5 - $15"],
            ["Azure Key Vault", "Standard (< 10K operations/mo)", "$0"],
            ["Azure Blob Storage", "Hot tier, < 1 GB", "< $1"],
            ["Infrastructure Total", "", "$5 - $16"],
            ["", "", ""],
            ["Power BI Pro", "$10/user/month per client", "$10 per client"],
            ["", "", ""],
            ["Total (1 client)", "", "~$15 - $26"],
            ["Total (5 clients)", "", "~$55 - $66"],
            ["Total (10 clients)", "", "~$105 - $116"],
        ]
    )

    doc.add_heading("Why Not Databricks?", level=2)
    doc.add_paragraph(
        "Databricks Standard starts at ~$100-200/month for always-on compute. At the scale "
        "of small business QBO data (1,000-10,000 transactions per month), you are paying "
        "for Spark cluster overhead that provides zero benefit. Azure Functions processes "
        "the same data in seconds, for free."
    )

    doc.add_page_break()

    # ================================================================
    # 14. TROUBLESHOOTING
    # ================================================================
    doc.add_heading("14. Troubleshooting", level=1)

    issues = [
        (
            "OAuth authorization fails or times out",
            "Ensure QBO_CLIENT_ID and QBO_CLIENT_SECRET in .env match your Intuit app. "
            "Check that the Redirect URI in your Intuit app settings matches "
            "QBO_REDIRECT_URI (default: http://localhost:8080/callback). "
            "Make sure no other process is using port 8080."
        ),
        (
            "401/403 errors when pulling QBO data",
            "The access token has expired. The pipeline auto-refreshes, but if the refresh "
            "token is also expired (101 days), you need to re-authorize: "
            "python -m auth.qbo_auth_flow --client-id <id>"
        ),
        (
            "429 Too Many Requests",
            "The rate limiter should prevent this, but if it happens, the retry logic "
            "will back off automatically. If persistent, check if another process is "
            "also calling the QBO API."
        ),
        (
            "Power BI shows no data",
            "Verify data exists in the database: open data/dev.db in a SQLite viewer "
            "and check that tables like fact_invoice have rows. Check that client_id "
            "values match between the data and any RLS filters."
        ),
        (
            "Pipeline fails for a specific table",
            "Some QBO tables (like Deposit, CreditMemo) may return 0 records if the "
            "company doesn't use those features. The pipeline logs a warning and continues. "
            "Check the logs for the specific error message."
        ),
        (
            "Azure Function doesn't trigger",
            "Check the function's timer trigger in the Azure Portal. Verify the CRON "
            "expression in function_app.py (0 0 11 * * * = daily at 11:00 UTC). "
            "Check Application Insights for error logs."
        ),
        (
            "Azure SQL connection fails from Power BI",
            "Ensure the Azure SQL firewall allows your IP address. In the Azure Portal, "
            "go to SQL Server > Networking > add your client IP. If using Power BI Service "
            "for scheduled refresh, install and configure the On-Premises Data Gateway."
        ),
    ]

    for title, fix in issues:
        doc.add_heading(title, level=3)
        doc.add_paragraph(fix)

    # ================================================================
    # SAVE
    # ================================================================
    output_path = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        "Bayshore_QBO_Pipeline_Guide.docx"
    )
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    build_document()
